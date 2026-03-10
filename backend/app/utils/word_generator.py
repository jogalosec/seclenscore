"""
Generador de documentos Word (.docx) con python-docx.
Reemplaza phpoffice/phpword del backend PHP.

Documentos soportados:
  - ECR (Evaluación de Cumplimiento de Requisitos)
  - ERS (Evaluación de Riesgo de Seguridad)
  - PAC (Plan de Acciones Correctivas)
"""
from __future__ import annotations

import io
from datetime import datetime
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn


# ---------------------------------------------------------------
# Helpers de estilo
# ---------------------------------------------------------------

def _set_cell_bg(cell, hex_color: str) -> None:
    """Establece el color de fondo de una celda de tabla."""
    from docx.oxml import parse_xml
    shading = parse_xml(
        f'<w:shd {qn("xmlns:w")}="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        f'w:fill="{hex_color}" w:color="auto" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def _bold_run(paragraph, text: str, size: int = 11, color: Optional[str] = None) -> None:
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    if color:
        r, g, b = int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16)
        run.font.color.rgb = RGBColor(r, g, b)


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _add_kv_table(doc: Document, datos: Dict[str, Any]) -> None:
    """Tabla de dos columnas clave → valor."""
    table = doc.add_table(rows=len(datos), cols=2)
    table.style = "Table Grid"
    for i, (k, v) in enumerate(datos.items()):
        row = table.rows[i]
        row.cells[0].text = str(k)
        _bold_run(row.cells[0].paragraphs[0].runs[0] if row.cells[0].paragraphs[0].runs else
                  row.cells[0].paragraphs[0].add_run(), "")
        row.cells[0].paragraphs[0].runs[0].bold = True if row.cells[0].paragraphs[0].runs else False
        row.cells[1].text = str(v) if v is not None else "—"
        _set_cell_bg(row.cells[0], "E8EAF6")


def _portada(doc: Document, titulo: str, activo: str, fecha: str) -> None:
    """Añade una portada estándar."""
    doc.add_picture  # no usamos imagen; dejamos espacio
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _bold_run(p, titulo, size=20, color="1A237E")

    doc.add_paragraph()
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.add_run(f"Activo: {activo}")

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.add_run(f"Fecha: {fecha}")

    doc.add_page_break()


# ---------------------------------------------------------------
# ECR — Evaluación de Cumplimiento de Requisitos
# ---------------------------------------------------------------

def generar_ecr(
    activo: Dict[str, Any],
    evaluacion: Dict[str, Any],
    preguntas: List[Dict[str, Any]],
    normativas: Optional[List[Dict[str, Any]]] = None,
    fecha: Optional[str] = None,
) -> bytes:
    """
    Genera un informe ECR en Word.

    Args:
        activo: dict con campos nombre, tipo, descripcion, etc.
        evaluacion: dict con meta_value (respuestas p1..pN) y fecha.
        preguntas: lista de preguntas con id, duda, nivel.
        normativas: lista de normativas relacionadas (opcional).
        fecha: fecha de generación (hoy si None).

    Returns:
        Bytes del documento .docx.
    """
    doc = Document()
    fecha_str = fecha or datetime.now().strftime("%d/%m/%Y")

    _portada(doc, "Evaluación de Cumplimiento de Requisitos (ECR)", activo.get("nombre", "—"), fecha_str)

    # 1. Datos del activo
    _add_heading(doc, "1. Información del activo", level=1)
    _add_kv_table(doc, {
        "Nombre": activo.get("nombre"),
        "Tipo": activo.get("tipo_nombre") or activo.get("tipo"),
        "Descripción": activo.get("descripcion"),
        "Estado": "Archivado" if activo.get("archivado") else "Activo",
        "Expuesto": "Sí" if activo.get("expuesto") else "No",
    })
    doc.add_paragraph()

    # 2. Resumen de la evaluación
    _add_heading(doc, "2. Resumen de la evaluación", level=1)
    respuestas: Dict[str, Any] = evaluacion.get("meta_value") or {}
    if isinstance(respuestas, str):
        import json
        try:
            respuestas = json.loads(respuestas)
        except ValueError:
            respuestas = {}

    total = len(respuestas)
    cumplidos = sum(1 for v in respuestas.values() if v and int(v) >= 3)
    no_cumplidos = total - cumplidos

    _add_kv_table(doc, {
        "Fecha de evaluación": str(evaluacion.get("fecha", fecha_str)),
        "Total de preguntas respondidas": total,
        "Cumplidos (valor ≥ 3)": cumplidos,
        "No cumplidos (valor < 3)": no_cumplidos,
        "Porcentaje de cumplimiento": f"{round(cumplidos / total * 100, 1)}%" if total > 0 else "N/A",
    })
    doc.add_paragraph()

    # 3. Normativas relacionadas
    if normativas:
        _add_heading(doc, "3. Normativas de referencia", level=1)
        for norm in normativas:
            doc.add_paragraph(f"• {norm.get('nombre', '—')} (v{norm.get('version', '')})", style="List Bullet")
        doc.add_paragraph()

    # 4. Detalle de respuestas
    _add_heading(doc, "4. Detalle de respuestas", level=1)
    if preguntas:
        table = doc.add_table(rows=1 + len(preguntas), cols=4)
        table.style = "Table Grid"
        headers = ["ID", "Pregunta", "Nivel", "Respuesta"]
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].bold = True
            _set_cell_bg(cell, "1A237E")
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

        for row_idx, pregunta in enumerate(preguntas, start=1):
            pid = pregunta.get("id", "")
            valor = respuestas.get(str(pid)) or respuestas.get(f"p{pid}", "—")
            row = table.rows[row_idx]
            row.cells[0].text = str(pid)
            row.cells[1].text = pregunta.get("duda", "—")
            row.cells[2].text = str(pregunta.get("nivel", "—"))
            row.cells[3].text = str(valor)
            if row_idx % 2 == 0:
                for cell in row.cells:
                    _set_cell_bg(cell, "F5F5F5")

    # Footer
    doc.add_paragraph()
    p = doc.add_paragraph(f"Documento generado automáticamente el {fecha_str} — SecLensCore")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(8)
    p.runs[0].font.color.rgb = RGBColor(150, 150, 150)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------
# ERS — Evaluación de Riesgo de Seguridad
# ---------------------------------------------------------------

def generar_ers(
    activo: Dict[str, Any],
    bia: Optional[Dict[str, Any]] = None,
    evaluaciones: Optional[List[Dict[str, Any]]] = None,
    fecha: Optional[str] = None,
) -> bytes:
    """
    Genera un informe ERS en Word con los resultados BIA y evaluaciones.

    Args:
        activo: dict con información del activo.
        bia: dict con meta_value (dimensiones BIA calculadas).
        evaluaciones: lista de evaluaciones de cumplimiento.
        fecha: fecha de generación.

    Returns:
        Bytes del documento .docx.
    """
    doc = Document()
    fecha_str = fecha or datetime.now().strftime("%d/%m/%Y")

    _portada(doc, "Evaluación de Riesgo de Seguridad (ERS)", activo.get("nombre", "—"), fecha_str)

    # 1. Información del activo
    _add_heading(doc, "1. Información del activo", level=1)
    _add_kv_table(doc, {
        "Nombre": activo.get("nombre"),
        "Tipo": activo.get("tipo_nombre") or activo.get("tipo"),
        "Descripción": activo.get("descripcion"),
    })
    doc.add_paragraph()

    # 2. Análisis de Impacto (BIA)
    _add_heading(doc, "2. Business Impact Analysis (BIA)", level=1)
    if bia and bia.get("meta_value"):
        import json
        meta = bia["meta_value"]
        if isinstance(meta, str):
            try:
                meta = json.loads(meta)
            except ValueError:
                meta = {}

        # Si meta contiene las dimensiones calculadas
        if isinstance(meta, dict) and "global" in meta:
            _add_kv_table(doc, {
                "Índice Global": meta.get("global", "—"),
                "Confidencialidad": meta.get("Con", {}).get("total", "—") if isinstance(meta.get("Con"), dict) else meta.get("Con", "—"),
                "Integridad": meta.get("Int", {}).get("total", "—") if isinstance(meta.get("Int"), dict) else meta.get("Int", "—"),
                "Disponibilidad": meta.get("Dis", {}).get("total", "—") if isinstance(meta.get("Dis"), dict) else meta.get("Dis", "—"),
            })
        else:
            doc.add_paragraph("BIA registrado pero pendiente de cálculo de dimensiones.")
    else:
        doc.add_paragraph("No se ha registrado BIA para este activo.")
    doc.add_paragraph()

    # 3. Historial de evaluaciones
    _add_heading(doc, "3. Historial de evaluaciones", level=1)
    if evaluaciones:
        table = doc.add_table(rows=1 + len(evaluaciones), cols=3)
        table.style = "Table Grid"
        for i, h in enumerate(["ID", "Tipo", "Fecha"]):
            cell = table.rows[0].cells[i]
            cell.text = h
            _set_cell_bg(cell, "1A237E")
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].bold = True
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

        for row_idx, ev in enumerate(evaluaciones, start=1):
            row = table.rows[row_idx]
            row.cells[0].text = str(ev.get("id", ""))
            row.cells[1].text = str(ev.get("meta_key", ""))
            row.cells[2].text = str(ev.get("fecha", ""))
    else:
        doc.add_paragraph("No se han registrado evaluaciones para este activo.")

    doc.add_paragraph()
    p = doc.add_paragraph(f"Documento generado automáticamente el {fecha_str} — SecLensCore")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(8)
    p.runs[0].font.color.rgb = RGBColor(150, 150, 150)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------
# PAC — Plan de Acciones Correctivas
# ---------------------------------------------------------------

def generar_pac(
    activo: Dict[str, Any],
    acciones: List[Dict[str, Any]],
    fecha: Optional[str] = None,
) -> bytes:
    """
    Genera un informe PAC en Word.

    Args:
        activo: dict con información del activo.
        acciones: lista de acciones correctivas con campos:
                  descripcion, responsable, prioridad, estado, fecha_limite.
        fecha: fecha de generación.

    Returns:
        Bytes del documento .docx.
    """
    doc = Document()
    fecha_str = fecha or datetime.now().strftime("%d/%m/%Y")

    _portada(doc, "Plan de Acciones Correctivas (PAC)", activo.get("nombre", "—"), fecha_str)

    # 1. Información del activo
    _add_heading(doc, "1. Información del activo", level=1)
    _add_kv_table(doc, {
        "Nombre": activo.get("nombre"),
        "Tipo": activo.get("tipo_nombre") or activo.get("tipo"),
    })
    doc.add_paragraph()

    # 2. Resumen ejecutivo
    _add_heading(doc, "2. Resumen ejecutivo", level=1)
    pendientes = sum(1 for a in acciones if str(a.get("estado", "")).lower() in ("pendiente", "abierto", "open"))
    completadas = len(acciones) - pendientes
    _add_kv_table(doc, {
        "Total de acciones": len(acciones),
        "Pendientes": pendientes,
        "Completadas": completadas,
        "Fecha del plan": fecha_str,
    })
    doc.add_paragraph()

    # 3. Detalle de acciones
    _add_heading(doc, "3. Detalle de acciones correctivas", level=1)
    if acciones:
        cols = ["#", "Descripción", "Responsable", "Prioridad", "Estado", "Fecha límite"]
        table = doc.add_table(rows=1 + len(acciones), cols=len(cols))
        table.style = "Table Grid"
        for i, h in enumerate(cols):
            cell = table.rows[0].cells[i]
            cell.text = h
            _set_cell_bg(cell, "1A237E")
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].bold = True
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

        PRIORIDAD_COLOR = {"alta": "FFCDD2", "media": "FFF9C4", "baja": "E8F5E9"}

        for row_idx, accion in enumerate(acciones, start=1):
            row = table.rows[row_idx]
            row.cells[0].text = str(row_idx)
            row.cells[1].text = str(accion.get("descripcion") or accion.get("meta_value") or "—")
            row.cells[2].text = str(accion.get("responsable", "—"))
            prioridad = str(accion.get("prioridad", "—")).lower()
            row.cells[3].text = prioridad.capitalize()
            row.cells[4].text = str(accion.get("estado", "Pendiente"))
            row.cells[5].text = str(accion.get("fecha_limite") or accion.get("fecha", "—"))

            bg = PRIORIDAD_COLOR.get(prioridad)
            if bg:
                _set_cell_bg(row.cells[3], bg)
    else:
        doc.add_paragraph("No se han registrado acciones correctivas.")

    doc.add_paragraph()
    p = doc.add_paragraph(f"Documento generado automáticamente el {fecha_str} — SecLensCore")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(8)
    p.runs[0].font.color.rgb = RGBColor(150, 150, 150)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
